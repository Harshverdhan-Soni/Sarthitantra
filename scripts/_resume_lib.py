"""Shared resume/cover layout helpers (run 22)."""
import pathlib
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x59, 0x59, 0x59)
FONT = 'Calibri'
CONTENT_W = 10224
BASE = pathlib.Path(__file__).parent.parent


def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Twips(12240)
    sec.page_height = Twips(15840)
    sec.left_margin = sec.right_margin = Twips(1008)
    sec.top_margin = sec.bottom_margin = Twips(864)
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    return doc


def set_font(run, size=9.5, bold=False, color=None, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_section_head(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(title.upper()), 10, bold=True, color=BLUE)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '1F4E79')
    pBdr.append(bot); pPr.append(pBdr)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    if bold_prefix:
        set_font(p.add_run(bold_prefix), bold=True)
    set_font(p.add_run(text))
    return p


def add_role_row(doc, title, org, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right'); tab.set(qn('w:pos'), str(CONTENT_W))
    tabs.append(tab); pPr.append(tabs)
    set_font(p.add_run(title), bold=True)
    set_font(p.add_run('  -  ' + org), color=GRAY)
    set_font(p.add_run('\t' + dates), color=GRAY)
    return p


def add_proj_head(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    set_font(p.add_run(text), bold=True, color=BLUE)
    return p


def add_header(doc, headline):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run('HARSHVERDHAN SONI'), size=16, bold=True, color=BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    set_font(p.add_run(headline), size=10, color=GRAY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run('Guwahati, India (Remote)  |  harshverdhansoni@gmail.com  |  +91-6900390339  |  '
                       'linkedin.com/in/harshverdhan-soni-680b2b31  |  github.com/Harshverdhan-Soni'),
             size=8.5, color=GRAY)


def add_education(doc, focus):
    add_section_head(doc, 'Education')
    add_role_row(doc, 'PhD Scholar (Pursuing) - AI/ML (Generative AI)', 'NIT Silchar', 'GPA 7.5')
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.15)
    set_font(p.add_run('Focus: ' + focus), color=GRAY, italic=True)
    add_role_row(doc, 'M.Tech - Computer Science & Engineering', 'NIT Calicut', '2017 · GPA 7.0')
    add_role_row(doc, 'MCA - Master of Computer Applications', 'Shri Sai Institute of Technology, Ratlam', '2013 · GPA 8.01')


def cover_para(doc, text, space_after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.12
    set_font(p.add_run(text), size=10.5)
    return p


def cover_line(doc, text, bold=False, size=10.5, color=None, space_after=2, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p
